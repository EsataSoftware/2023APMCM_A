import os
import nbformat
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def extract_code_from_notebooks(directory_path, output_path):
    # 创建一个新的Word文档
    doc = Document()

    # 设置标题
    title = doc.add_heading('Jupyter Notebook Code Extract', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 遍历目录中的所有.ipynb文件
    for filename in os.listdir(directory_path):
        if filename.endswith('.ipynb'):
            notebook_path = os.path.join(directory_path, filename)
            
            # 读取Jupyter笔记本
            with open(notebook_path, 'r', encoding='utf-8') as nb_file:
                notebook_content = nbformat.read(nb_file, as_version=4)

            # 提取代码块
            for cell in notebook_content['cells']:
                if cell['cell_type'] == 'code':
                    code_source = cell['source']

                    # 在文档中添加新的代码块
                    doc.add_paragraph('Code from: ' + filename, style='Heading2')
                    code_paragraph = doc.add_paragraph(code_source)
                    code_paragraph.style.font.size = Pt(10)

    # 保存Word文档
    doc.save(output_path)

if __name__ == "__main__":
    input_directory = 'C:\\Users\\Lenovo\\Desktop\\Yolov5Test\\Aapmcm2308882fj\\'
    output_file = 'code_extraction.docx'
    extract_code_from_notebooks(input_directory, output_file)
